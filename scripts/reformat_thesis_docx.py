from __future__ import annotations

import re
import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
REPORTS = ROOT / "reports"
THESIS_TITLE = "新能源储能侧优化调度系统设计与实现"
POLLUTED_STRINGS = [
    "基于智能技术的电力变压器故障诊断系统",
    "电力变压器",
    "Delphi7",
    "Access 数据库",
    "提示：",
    "TODO",
    "待补",
]


def set_run_font(run, name: str, size_pt: float, bold: bool = False) -> None:
    """Apply both Western and East Asian font names so Word renders Chinese predictably."""
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size_pt)
    run.bold = bold


def set_paragraph_spacing(paragraph, line_spacing: float = 1.5, first_line_chars: int = 0) -> None:
    """Set common thesis paragraph spacing; first-line indent uses the common 2-character width."""
    fmt = paragraph.paragraph_format
    fmt.line_spacing = line_spacing
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    if first_line_chars:
        fmt.first_line_indent = Pt(10.5 * first_line_chars)


def set_paragraph_bottom_border(paragraph) -> None:
    """Add the thin horizontal line used by the sample header."""
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    p_bdr.append(bottom)


def add_field(paragraph, instr: str) -> None:
    """Insert a Word field such as PAGE or TOC without relying on pre-rendered text."""
    def append_run_child(child):
        run = OxmlElement("w:r")
        run.append(child)
        paragraph._p.append(run)

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    append_run_child(begin)

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instr
    append_run_child(instr_text)

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    append_run_child(separate)

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    append_run_child(end)


def restart_page_numbering(section, start: int = 1) -> None:
    """Restart page numbering for the section that begins the thesis body."""
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn("w:start"), str(start))


def configure_section(section, with_footer: bool) -> None:
    """Apply school-sample-like page margins, header title, and optional body page number."""
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(2.50)
    section.header_distance = Cm(1.50)
    section.footer_distance = Cm(1.75)

    header = section.header
    header.is_linked_to_previous = False
    paragraph = header.paragraphs[0]
    paragraph.text = ""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(THESIS_TITLE)
    set_run_font(run, "宋体", 9)
    set_paragraph_bottom_border(paragraph)

    footer = section.footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0]
    paragraph.text = ""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if with_footer:
        run = paragraph.add_run("- ")
        set_run_font(run, "Times New Roman", 9)
        add_field(paragraph, " PAGE ")
        run = paragraph.add_run(" -")
        set_run_font(run, "Times New Roman", 9)


def configure_styles(doc: Document) -> None:
    """Normalize built-in styles so automatic TOC and headings remain Word-compatible."""
    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)

    for name in ("Heading 1", "Heading 2", "Heading 3"):
        style = doc.styles[name]
        style.font.name = "宋体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.font.color.rgb = None
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.5

    doc.styles["Heading 1"].font.size = Pt(15)
    doc.styles["Heading 1"].font.bold = True
    doc.styles["Heading 2"].font.size = Pt(12)
    doc.styles["Heading 2"].font.bold = True
    doc.styles["Heading 3"].font.size = Pt(12)
    doc.styles["Heading 3"].font.bold = True


def add_center_title(doc: Document, text: str, font_name: str, size_pt: float, bold: bool = True) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(paragraph, 1.5)
    run = paragraph.add_run(text)
    set_run_font(run, font_name, size_pt, bold)


def add_body_paragraph(doc: Document, text: str, english: bool = False) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_spacing(paragraph, 1.5, 2 if not english else 0)
    run = paragraph.add_run(text)
    set_run_font(run, "Times New Roman" if english else "宋体", 12)


def add_keyword_paragraph(doc: Document, text: str, english: bool = False) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(paragraph, 1.5)
    label = "Key Words:" if english else "关键词："
    if text.startswith(label):
        value = text[len(label) :].strip()
    else:
        value = text
    # Chinese sample keeps the value immediately after "关键词："; English keeps a space
    # after "Key Words:" according to conventional English typography.
    r1 = paragraph.add_run(label + (" " if english else ""))
    set_run_font(r1, "Times New Roman" if english else "黑体", 12, True)
    r2 = paragraph.add_run(value)
    set_run_font(r2, "Times New Roman" if english else "宋体", 12)


def add_toc(doc: Document) -> None:
    add_center_title(doc, "目录", "黑体", 15, True)
    paragraph = doc.add_paragraph()
    # Word updates this field into a real automatic TOC: levels 1-3, hyperlinks enabled.
    add_field(paragraph, r'TOC \o "1-3" \h \z \u')


def add_heading(doc: Document, text: str) -> None:
    if re.match(r"^第.+章\s+", text):
        paragraph = doc.add_paragraph(style="Heading 1")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(paragraph, 1.5)
        run = paragraph.add_run(text)
        set_run_font(run, "黑体", 15, True)
        return
    if re.match(r"^\d+\.\d+\.\d+\s+", text):
        paragraph = doc.add_paragraph(style="Heading 3")
    else:
        paragraph = doc.add_paragraph(style="Heading 2")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(paragraph, 1.5)
    run = paragraph.add_run(text)
    set_run_font(run, "黑体", 12, True)


def add_source_code_block(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(paragraph, 1.0)
    run = paragraph.add_run(text)
    set_run_font(run, "Consolas", 9)


def set_update_fields_on_open(doc: Document) -> None:
    """Ask Word to refresh fields such as TOC and PAGE when the document is opened."""
    settings = doc.settings._element
    existing = settings.find(qn("w:updateFields"))
    if existing is None:
        existing = OxmlElement("w:updateFields")
        settings.append(existing)
    existing.set(qn("w:val"), "true")


def classify_source(source: Path):
    doc = Document(str(source))
    paragraphs = [(p.text.strip(), p.style.name) for p in doc.paragraphs if p.text.strip()]

    first_chapter_idx = next(i for i, (t, _) in enumerate(paragraphs) if t.startswith("第一章"))
    abstract_idx = next(i for i, (t, _) in enumerate(paragraphs) if t == "摘 要")
    english_title_idx = next(i for i, (t, _) in enumerate(paragraphs) if t.startswith("Design and Implementation"))
    english_abstract_idx = next(i for i, (t, _) in enumerate(paragraphs) if t == "Abstract")

    cn_abstract = paragraphs[abstract_idx + 1 : english_title_idx]
    cn_body = [t for t, _ in cn_abstract if not t.startswith("关键词")]
    cn_keywords = next((t for t, _ in cn_abstract if t.startswith("关键词")), "")

    en_abstract = paragraphs[english_abstract_idx + 1 : first_chapter_idx]
    en_body = [t for t, _ in en_abstract if not t.startswith("Key Words")]
    en_keywords = next((t for t, _ in en_abstract if t.startswith("Key Words")), "")

    body = paragraphs[first_chapter_idx:]
    return {
        "cn_body": cn_body,
        "cn_keywords": cn_keywords,
        "en_title": paragraphs[english_title_idx][0],
        "en_body": en_body,
        "en_keywords": en_keywords,
        "body": body,
        "paragraph_count": len(paragraphs),
    }


def locate_source() -> Path:
    candidates = [
        p
        for p in REPORTS.glob("*.docx")
        if not p.name.startswith(("~$", "14.", "6."))
        and "格式修正版" not in p.name
        and "备份" not in p.name
    ]
    if not candidates:
        raise FileNotFoundError("未找到可用的初稿 DOCX。")
    return max(candidates, key=lambda path: path.stat().st_mtime)


def build_document(source: Path, output: Path) -> dict:
    data = classify_source(source)
    doc = Document()
    configure_styles(doc)
    configure_section(doc.sections[0], with_footer=False)

    add_center_title(doc, "摘 要", "黑体", 15, True)
    for text in data["cn_body"]:
        add_body_paragraph(doc, text, english=False)
    add_keyword_paragraph(doc, data["cn_keywords"], english=False)
    doc.add_page_break()

    add_center_title(doc, data["en_title"], "Times New Roman", 18, True)
    add_center_title(doc, "Abstract", "Times New Roman", 15, True)
    for text in data["en_body"]:
        add_body_paragraph(doc, text, english=True)
    add_keyword_paragraph(doc, data["en_keywords"], english=True)
    doc.add_page_break()

    add_toc(doc)

    body_section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(body_section, with_footer=True)
    restart_page_numbering(body_section, 1)

    first_chapter_seen = False
    for text, style_name in data["body"]:
        if re.match(r"^第.+章\s+", text):
            # The school sample starts every chapter on a new page. The first chapter
            # already starts a new body section, so only later chapters need page breaks.
            if first_chapter_seen:
                doc.add_page_break()
            first_chapter_seen = True
        if style_name.startswith("Heading") or re.match(r"^第.+章\s+", text) or re.match(r"^\d+\.\d+", text):
            add_heading(doc, text)
        elif style_name == "Source Code":
            add_source_code_block(doc, text)
        elif re.match(r"^式（\d+-\d+）", text) or re.match(r"^[A-Za-z_].*[=<>]", text):
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_spacing(paragraph, 1.5)
            run = paragraph.add_run(text)
            set_run_font(run, "Times New Roman", 12)
        else:
            add_body_paragraph(doc, text, english=False)

    set_update_fields_on_open(doc)
    doc.save(str(output))
    return data


def extract_docx_text(path: Path) -> str:
    doc = Document(str(path))
    texts = [p.text for p in doc.paragraphs]
    for section in doc.sections:
        texts.extend(p.text for p in section.header.paragraphs)
        texts.extend(p.text for p in section.footer.paragraphs)
    return "\n".join(texts)


def write_report(report: Path, source: Path, backup: Path, output: Path, data: dict, checks: dict) -> None:
    status = "通过" if all(checks.values()) else "存在需处理项"
    lines = [
        "# 初稿 DOCX 格式修复报告",
        "",
        f"- 源文件：`{source}`",
        f"- 备份文件：`{backup}`",
        f"- 输出文件：`{output}`",
        f"- 修复时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        f"- 自动检查结论：{status}",
        "",
        "## 已修复项",
        "- 删除初稿开头的独立论文标题页式段落，正文前置顺序改为中文摘要、英文摘要、目录、第一章正文。",
        "- 新建自动目录域：`TOC \\o \"1-3\" \\h \\z \\u`，供 Word 执行“更新整个目录”。",
        "- 页眉统一改为论文题名：新能源储能侧优化调度系统设计与实现。",
        "- 正文从第一章开始新分节，并将页码重置为 `- 1 -` 格式。",
        "- 重建摘要、英文摘要、一级标题、二级标题、三级标题、正文段落、公式/代码段的基础样式。",
        "- 输出过程中未修改原始初稿 DOCX。",
        "",
        "## 自动检查结果",
    ]
    for name, passed in checks.items():
        lines.append(f"- {'通过' if passed else '未通过'}：{name}")
    lines.extend(
        [
            "",
            "## 需要 Word 人工完成的项",
            "- 打开格式修正版后，右键目录并选择“更新域/更新整个目录”。",
            "- 检查摘要页、英文摘要页、目录页、第一章第一页、参考文献页的视觉分页是否与范本完全一致。",
            "- 若学校要求封面、诚信声明或任务书一起提交，需要在本文件之前另行合并学校指定封面材料。",
            "",
            "## Pitfall",
            "- Word 自动目录域在脚本生成后不会自动渲染页码；未更新目录前看到空目录或提示文字是正常现象，不代表目录结构失败。",
            "",
            "## 阶段进度",
            "- 本阶段目标：修复初稿 DOCX 的学校模板结构和主要排版污染。",
            "- 完成情况：已生成修正版、备份和修复报告；剩余工作是 Word 更新目录和人工视觉核验。",
            "- 下一阶段可行性：高。若视觉核验发现局部字号或分页偏差，可在当前脚本基础上继续微调样式。",
        ]
    )
    report.write_text("\n".join(lines), encoding="utf-8")


def main() -> None:
    source = locate_source()
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup = source.with_name(f"{source.stem}_格式修复前备份_{timestamp}{source.suffix}")
    output = source.with_name(f"{source.stem}_格式修正版{source.suffix}")
    report = source.with_name(f"{source.stem}_格式修复报告.md")

    shutil.copy2(source, backup)
    data = build_document(source, output)

    text = extract_docx_text(output)
    order_tokens = ["摘 要", "Abstract", "目录", "第一章"]
    positions = [text.find(token) for token in order_tokens]
    checks = {
        "前置顺序为中文摘要、英文摘要、目录、第一章": all(pos >= 0 for pos in positions)
        and positions == sorted(positions),
        "不存在范本污染文字": not any(s in text for s in POLLUTED_STRINGS),
        "页眉包含新论文题名": THESIS_TITLE in text,
        "参考文献保留": "参考文献" in text,
        "正文段落已读取": data["paragraph_count"] > 100,
    }
    write_report(report, source, backup, output, data, checks)

    print(f"SOURCE={source}")
    print(f"BACKUP={backup}")
    print(f"OUTPUT={output}")
    print(f"REPORT={report}")
    for key, value in checks.items():
        print(f"CHECK {key}: {'PASS' if value else 'FAIL'}")
    if not all(checks.values()):
        raise SystemExit(2)


if __name__ == "__main__":
    main()
