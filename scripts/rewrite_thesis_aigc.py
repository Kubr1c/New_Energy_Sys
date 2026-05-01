from __future__ import annotations

import re
import shutil
from collections import Counter
from datetime import datetime
from pathlib import Path

from docx import Document
from lxml import html


ROOT = Path(__file__).resolve().parents[1]
REPORTS = ROOT / "reports"


def compact(text: str) -> str:
    """Collapse Word/PaperYY whitespace so long Chinese paragraphs can be matched safely."""
    return " ".join(text.split())


def find_source_docx() -> Path:
    """Find the current thesis draft while ignoring school templates and generated outputs."""
    candidates = [
        path
        for path in REPORTS.glob("*.docx")
        if not path.name.startswith(("~$", "14.", "6."))
        and "降AIGC" not in path.name
        and "格式修复前备份" not in path.name
    ]
    if not candidates:
        raise FileNotFoundError("未找到可用于降 AIGC 改写的论文 DOCX。")
    return max(candidates, key=lambda path: path.stat().st_mtime)


def find_report_html() -> Path:
    """Find the PaperYY AIGC HTML report supplied by the user."""
    candidates = [path for path in REPORTS.glob("*.html") if "AIGC" in path.name]
    if not candidates:
        raise FileNotFoundError("未找到 AIGC 检测报告 HTML。")
    return max(candidates, key=lambda path: path.stat().st_mtime)


def parse_hotspots(report_path: Path) -> list[tuple[str, str]]:
    """Extract high/medium sentence fragments from the report for traceability."""
    tree = html.fromstring(report_path.read_text(encoding="utf-8", errors="ignore"))
    hotspots: list[tuple[str, str]] = []
    for em in tree.xpath("//em"):
        cls = em.get("class") or ""
        if cls in {"high", "medium", "middle"}:
            text = compact("".join(em.itertext()))
            if len(text) >= 8:
                hotspots.append((cls, text))
    return hotspots


# The replacement text below is intentionally project-grounded rather than a synonym rewrite.
# It preserves reported metrics, data sources, model names, and experimental boundaries while
# reducing template-like "负责/完成/表明/验证" prose that PaperYY marked as high risk.
REPLACEMENTS: dict[str, str] = {
    "系统以 PVDAQ 光伏实测数据、NSRDB 太阳资源数据和 OPSD 负荷、电价画像为基础，构建了数据采集、清洗、特征工程、模型训练、主模型推理、储能调度、策略治理和前端展示的完整流程。在预测层，本文对 LightGBM、XGBoost、CatBoost、随机森林、TCN、CNN-LSTM 和 Attention-LSTM 等模型进行了对比。实验结果表明，在仅使用预测时刻之前可获得信息的 history_only 输入条件下，调优后的 LightGBM 在 t+24h 光伏功率预测任务上取得测试集 nRMSE 0.1225、日间 nRMSE 0.1689，综合稳定性和可解释性较好，因此被固化为系统主预测模型。深度学习模型在目标时刻太阳资源特征参与的离线上限实验中表现更优，但在真实可获得输入条件下尚未达到替代 LightGBM 的工程阈值。":
        "本文的数据链路从 PVDAQ 光伏实测功率开始，按小时粒度对齐 NSRDB 太阳资源字段，并引入 OPSD 负荷与电价画像作为调度侧输入。围绕 t+24h 预测任务，项目分别训练 LightGBM、XGBoost、CatBoost、随机森林、TCN、CNN-LSTM 和 Attention-LSTM，并把可获得输入边界单独列为模型筛选条件。最终保留的是 tuned LightGBM history_only 模型：该模型在测试集上的 nRMSE 为 0.1225，日间 nRMSE 为 0.1689。CNN-LSTM 在带有目标时刻太阳资源字段的离线上限实验中误差更低，但这类字段不能作为真实预测时刻的输入，因此本文没有把深度学习结果直接写成主模型替换结论。",
    "在调度层，本文以主模型预测结果为输入，设计了固定阈值调度、分位数阈值扫描、24h 滚动优化调度和储能配置敏感性分析方法。实验表明，固定阈值策略因放电阈值超出样本电价分布而产生负增量收益；分位数阈值策略可作为离线上界；滚动优化策略在约束可审计条件下实现正增量收益；储能容量、功率和目标函数惩罚项对收益、循环次数、短缺量和 SOC 贴边风险具有明显影响。系统后端采用 FastAPI 实现数据与任务接口，前端采用 Vue 3、Element Plus 和 ECharts 实现预测、调度、治理和报告展示。本文工作验证了公开数据条件下新能源储能侧“预测-调度-治理-展示”闭环的可行性，并明确了离线收益、真实市场价格和 forecast-cycle 天气数据之间的边界。":
        "储能调度部分直接使用主模型预测产物，先用固定阈值策略暴露参数与电价分布不匹配的问题，再用分位数阈值扫描给出离线对照，随后通过 24h rolling look-ahead 方案检查 SOC、功率、同时充放电和能量守恒等约束。Stage12 的滚动优化能够取得正增量收益，但收益不如 Stage11 的离线上界；Stage15 进一步说明容量、功率倍率和目标函数惩罚项会同时改变收益、循环次数、短缺量和 SOC 贴边比例。系统展示层没有直接读取零散 CSV，而是由 FastAPI 汇总阶段产物，再交给 Vue 3、Element Plus 和 ECharts 页面展示。本文的结论因此限定在公开数据回放场景内，OPSD 映射电价和 forecast-cycle 天气输入都没有被写成真实生产结算能力。",
    "The system uses PVDAQ photovoltaic measurements, NSRDB solar-resource data and OPSD-based load and price profiles to construct a complete pipeline covering data acquisition, cleaning, feature engineering, model training, main-model inference, storage dispatch, strategy governance and frontend visualization. In the forecasting layer, LightGBM, XGBoost, CatBoost, random forests, TCN, CNN-LSTM and Attention-LSTM are compared. Experimental results show that the tuned LightGBM model under the history_only feature setting achieves a test nRMSE of 0.1225 and a daytime nRMSE of 0.1689 for the t+24h photovoltaic forecasting task. Considering stability, interpretability and production-safe input boundaries, LightGBM is selected as the main forecasting model. Deep learning models show stronger performance in offline upper-bound experiments with target-time solar-resource features, but they do not yet satisfy the replacement threshold under strictly available historical inputs.":
        "The data workflow starts from PVDAQ photovoltaic measurements, aligns them with NSRDB solar-resource variables, and then adds OPSD-based load and price profiles for the dispatch experiment. For the t+24h forecasting task, the project compares LightGBM, XGBoost, CatBoost, random forests, TCN, CNN-LSTM and Attention-LSTM under the same evaluation protocol. The tuned LightGBM model with the history_only feature set is retained as the main model because it reaches a test nRMSE of 0.1225 and a daytime nRMSE of 0.1689 without using target-time information. CNN-LSTM performs better only in the offline upper-bound setting where target-aligned solar-resource fields are available, so this result is treated as a boundary analysis rather than a production-safe replacement.",
    "In the dispatch layer, the thesis designs fixed-threshold dispatch, quantile-threshold scanning, 24-hour rolling look-ahead dispatch and storage-configuration sensitivity analysis. Results show that the fixed-threshold strategy fails because its discharge threshold is outside the observed price range, while the quantile strategy provides an offline upper bound and the rolling strategy achieves positive incremental revenue under auditable constraints. Storage capacity, power limits and objective penalties significantly affect revenue, cycle count, shortfall and SOC boundary risk. The backend is implemented with FastAPI, and the frontend is implemented with Vue 3, Element Plus and ECharts. This work demonstrates the feasibility of a renewable energy storage-side forecasting-dispatch-governance-visualization loop under public-data conditions, while explicitly clarifying the boundaries of offline revenue, real market prices and forecast-cycle weather data.":
        "On the dispatch side, the main-model forecast is first tested with a fixed threshold rule, then compared with quantile scanning and a 24-hour rolling look-ahead strategy. The fixed threshold rule produces negative incremental revenue because the discharge threshold is higher than the observed price range. The quantile rule is therefore used only as an offline upper-bound reference, whereas the rolling strategy is kept as the auditable pilot candidate. Stage15 also shows that storage capacity, power limits and penalty terms change revenue, cycle count, shortfall and SOC boundary behavior at the same time. FastAPI is used to expose these stage outputs to the Vue 3 frontend, where Element Plus and ECharts present forecasting, dispatch, governance and report views. The reported revenue remains an offline OPSD-mapped result rather than a real Colorado market settlement.",
    "随着能源结构转型持续推进，光伏、风电等新能源发电在电力系统中的占比不断提高。与传统火电、水电等可控电源相比，新能源发电具有更强的随机性和不确定性。以光伏发电为例，其出力受太阳辐照度、云量、温度、季节和昼夜周期共同影响，短时间内可能出现明显波动。当新能源出力预测误差较大或调节能力不足时，电力系统可能面临弃光、功率爬坡、供需不平衡和调度成本增加等问题。":
        "在以光伏为代表的新能源接入场景中，调度系统面对的不是稳定可控的机组出力，而是随辐照度、云量、温度和昼夜周期不断变化的功率序列。项目选用的 PVDAQ system 10 数据也体现了这种特点：同一电站在不同季节和天气条件下的小时级功率差异明显。若日前预测偏差较大，后续储能策略会在错误的功率预期上安排充放电，进而放大短缺、弃光或收益波动。因此，本文把光伏预测和储能调度放在同一条实验链路中处理，而不是只讨论单个模型的误差。",
    "储能系统能够在时间维度上转移电能，是支撑新能源消纳和灵活调节的重要资源。对于新能源侧场景，储能既可以吸纳光伏高出力时段的富余电量，也可以在光伏不足、负荷较高或电价较高时放电，从而实现平滑出力、降低短缺风险和提高经济收益等目标。但储能调度并不是简单的“低价充电、高价放电”。实际调度需要同时考虑功率上限、容量约束、SOC 安全边界、充放电效率、循环损耗、电价分布、预测误差和终端 SOC 等因素。如果调度策略参数与数据分布不匹配，储能可能出现只充不放、频繁贴边、收益下降或循环次数过高等问题。":
        "储能在本文中被建模为带容量、功率和 SOC 边界的调节资源。它可以把光伏高出力时段的部分电量转移到价格或负荷更高的时段，但调度结果取决于一组相互牵制的条件：充放电效率会改变可用能量，循环成本会压低频繁动作的收益，终端 SOC 惩罚又会影响滚动窗口最后几步的决策。项目中的 Stage10 固定阈值实验说明，若放电阈值没有落在样本电价分布内，储能会出现充放电行为失衡，甚至产生负增量收益。这个现象也是后续引入分位数阈值和滚动优化的直接原因。",
    "近年来，机器学习和深度学习方法被广泛用于新能源功率预测与电力系统调度。LightGBM、XGBoost 等表格模型具有训练速度快、适合结构化特征、可解释性较强等优点；TCN、LSTM、Transformer 等深度学习模型在时序依赖建模方面具有表达能力；强化学习和滚动优化方法则常用于储能调度和综合能源系统能量管理。已有研究为本文提供了理论基础，但在本科毕业设计场景中，仍需要解决一个更工程化的问题：如何在缺少真实电站私有数据和真实市场结算数据的情况下，利用公开数据构建一个可复现、可解释、可展示的新能源储能侧调度系统。":
        "已有研究通常从预测模型精度、储能调度算法或源网荷储协同控制中的某一方面展开。本文借鉴这些方法，但约束条件更具体：实验数据来自公开渠道，电价并非真实同区域结算价格，天气数据也要区分历史回放和真实 forecast-cycle 输入。在这样的条件下，直接追求复杂模型并不一定能带来可交付结果。本文更关注一条可复现的工程路径，即先用可审计的特征和模型得到日前光伏预测，再让储能策略在明确的 SOC、功率和收益口径下运行，最后通过前后端系统把各阶段结果展示出来。",
    "本文围绕“光伏功率预测 + 储能调度优化 + 系统可视化展示”展开研究。系统采用公开光伏、太阳资源、负荷和电价画像数据，重点建立从数据处理到预测模型、从预测结果到储能调度、从调度指标到前端展示的闭环。与单纯模型实验不同，本文强调模型输入边界、离线实验口径、调度约束审计和系统交付能力，避免将离线上限结果误写为真实上线能力。":
        "围绕上述问题，本文把工作拆成三个可检查的部分：第一部分处理 PVDAQ、NSRDB 和 OPSD 数据，使预测和调度使用同一时间轴；第二部分比较表格模型与深度学习模型，并把 history_only 作为主模型输入边界；第三部分把 Stage9 预测结果送入储能调度、治理评分和前端页面。这样安排的重点不在于给出一个脱离场景的最高分模型，而是让每个结论都能回到数据来源、阶段报告和系统接口中核验。",
    "新能源功率预测研究主要经历了统计模型、传统机器学习和深度学习三个阶段。早期方法包括 ARIMA、回归模型和支持向量机等，这类方法实现简单，但对非线性和复杂气象因素的表达能力有限。随后，随机森林、梯度提升树、LightGBM、XGBoost 等模型被广泛应用于光伏和风电预测任务，能够较好处理多源结构化特征。近年来，LSTM、GRU、TCN、Transformer 及其改进模型被用于捕捉长短期时间依赖关系，在部分数据充分、输入特征质量较高的任务中表现较好[1-3]。":
        "新能源功率预测方法大致可以从输入数据和模型结构两个角度理解。ARIMA、回归模型等方法对数据要求较低，但很难充分表达辐照度、温度、历史功率和时间周期之间的非线性关系。随机森林、LightGBM、XGBoost 等树模型更适合本文这类结构化特征表，训练过程也便于复现。LSTM、TCN、Transformer 等深度模型能够处理序列依赖，但通常更依赖样本规模、窗口构造和特征可获得性。本文在文献基础上保留多类模型对比，是为了判断公开数据条件下哪类方法更适合作为系统主链路[1-3]。",
    "储能调度研究主要包括规则策略、数学优化、模型预测控制和强化学习等方向。规则策略实现简单、可解释性强，但依赖人工阈值，容易受到电价分布和负荷场景变化影响。数学优化方法可显式建模 SOC、功率和收益目标，但当系统规模增大或目标函数复杂时，计算成本和参数设计难度会上升。滚动优化方法在每个时刻利用未来窗口信息生成计划，并只执行当前动作，兼具可解释性和实时修正能力。深度强化学习方法能够处理连续动作和复杂状态，但训练稳定性、约束安全性和可解释性仍是工程应用中的主要挑战[4-6]。":
        "储能调度方法的差异主要体现在约束表达和决策方式上。规则策略只需要少量阈值，便于解释，但 Stage10 的结果表明，阈值一旦偏离电价分布，策略会很快失效。数学优化可以把 SOC、功率、收益和惩罚项写入目标函数，适合做离线分析。滚动优化每次只执行当前时刻动作，下一时刻再依据新状态重新求解，因而比一次性离线策略更接近实际调度。强化学习方法适合复杂状态空间，但对仿真环境、训练稳定性和约束处理要求较高，本文只在文献综述中讨论其适用性[4-6]。",
    "在源网荷储协同和综合能源系统研究中，多智能体强化学习、共享储能、虚拟电厂和光储充一体化等主题受到关注。相关研究表明，深度强化学习适合探索复杂调度策略，但其有效性依赖高质量仿真环境和严格的安全约束机制。对于本文这样的本科工程原型，若直接将强化学习作为主调度算法，可能面临训练成本高、结果难解释、约束难审计等风险。因此，本文选择以可审计的 LightGBM 预测模型和滚动优化调度作为主线，同时保留深度学习预测模型作为对比验证。":
        "源网荷储协同研究为本文提供了方法背景，但本文没有把多智能体强化学习作为主算法。原因在于，本项目首先要保证公开数据回放、预测输入边界和储能约束检查都能被复现；如果在此基础尚未完全稳定时引入强化学习，训练环境、奖励函数和安全约束会成为新的不确定来源。因此，本文把 LightGBM history_only 预测和滚动优化调度作为主线，把深度学习预测模型放在对比实验中，用来观察复杂模型在不同输入边界下的收益和限制。",
    "第一，构建公开数据驱动的新能源预测与调度数据集。系统接入 PVDAQ 光伏实测数据、NSRDB 太阳资源数据和 OPSD 负荷、电价画像数据，完成时间对齐、缺失值处理、异常值检查和小时级特征工程。":
        "第一，整理公开数据并形成预测、调度共用的数据底座。项目以 PVDAQ system 10 的实测功率为主线，将 NSRDB PSM 太阳资源字段和 OPSD 负荷、电价画像对齐到小时级时间戳，同时保留缺失值、异常值、重复时间和单位转换检查记录。",
    "本文共分为七章。第一章介绍研究背景、研究现状和研究内容。第二章介绍光伏功率预测、储能模型、滚动优化和评价指标等理论基础。第三章给出系统需求分析与总体设计。第四章介绍数据处理和光伏功率预测模型实现。第五章介绍储能调度策略、治理评分和敏感性分析。第六章介绍后端 API、前端展示和系统测试。第七章总结全文工作并提出后续改进方向。":
        "全文结构按照项目实现顺序展开。第一章先说明新能源侧预测与储能调度为什么需要放在同一系统中讨论；第二章交代后续会用到的预测模型、储能约束和评价指标；第三章从需求和数据流角度描述系统总体设计；第四章对应数据处理和预测模型选择；第五章转入储能调度、治理评分和配置敏感性分析；第六章说明 FastAPI 后端、Vue 前端以及测试边界；第七章回到已完成工作和仍需改进的部分。",
    "为避免时间泄漏，本文在主模型中采用 history_only 特征组，即只使用预测时刻之前可获得的信息。部分实验中使用与目标时刻对齐的太阳资源字段作为离线上限分析，但该类特征不能被描述为真实上线输入。这个边界是本文实验设计的关键。":
        "本文把特征可获得性作为模型实验的前置约束。主模型只使用 history_only 特征，也就是预测时刻之前已经出现的功率、天气和时间信息。另一些实验引入目标时刻对齐的太阳资源字段，目的只是估计离线上限：它们能够帮助判断模型潜力，却不能被写入真实预测流程。后续所有模型选择都以这个边界为准。",
    "LightGBM 是一种基于梯度提升决策树的机器学习模型。它通过逐步构建多棵弱学习器，最小化损失函数，从而得到具有较强非线性表达能力的集成模型。与传统 GBDT 相比，LightGBM 在训练速度、内存占用和大规模特征处理方面具有优势，适合结构化数据场景。":
        "LightGBM 属于梯度提升树模型，其基本思路是连续训练多棵决策树，让后续树修正前面模型留下的误差。对本文而言，LightGBM 的价值不只在模型结构本身，更在于它能直接处理 Stage3 生成的结构化特征表，包括历史功率、时间特征、天气字段、负荷和电价画像。相比需要复杂窗口输入的深度模型，它的训练、推理和特征检查更容易嵌入后端流程。",
    "本文选择 LightGBM 作为工程主模型，主要原因包括：第一，项目数据由时间、历史功率、天气、负荷、电价等结构化字段组成，适合树模型处理；第二，LightGBM 训练和推理稳定，便于在后端服务中固化；第三，模型输入特征可控，方便检查是否包含未来信息；第四，在本项目同口径对比中，LightGBM 在 history_only 条件下综合表现最好。":
        "主模型最终采用 LightGBM，是由数据形态和实验门禁共同决定的。Stage3 输出的是以小时为单位的结构化特征表，树模型可以直接利用这些字段；Stage9 推理链路又要求模型 bundle、特征列、时间顺序和物理裁剪规则都能被固定下来。在同口径比较中，LightGBM tuned history_only 的整体 nRMSE 和日间 nRMSE 达到当前主线要求，并且没有依赖目标时刻后验特征，因此比复杂模型更适合作为系统默认预测入口。",
    "深度学习模型能够自动学习时序数据中的非线性关系。本文实现了 TCN、CNN-LSTM 和 Attention-LSTM 等模型作为对比。TCN 使用一维卷积和扩张卷积捕捉历史窗口内的时间依赖；CNN-LSTM 先通过卷积提取局部模式，再通过 LSTM 建模序列依赖；Attention-LSTM 在 LSTM 基础上引入注意力机制，尝试提高模型对关键时刻的关注能力。":
        "深度学习实验主要用于检查历史窗口建模是否能带来额外收益。本文实现的 TCN 通过一维卷积和扩张卷积处理固定长度序列；CNN-LSTM 先抽取局部变化模式，再把结果交给 LSTM；Attention-LSTM 则在序列输出上加入注意力权重。它们与 LightGBM 使用同一预测任务，但输入组织方式不同，因此更适合作为模型能力对照，而不是默认替代项。",
    "储能调度必须满足以下约束：":
        "在本文的离线回放中，储能动作先接受以下约束检查：",
    "同一时刻不能同时充电和放电。实际调度还需考虑循环成本、短缺惩罚、弃光惩罚和终端 SOC 目标。本文在各阶段报告中对 SOC 越界、功率越限、同时充放电和能量守恒误差进行了质量门禁检查。":
        "同一时间步只允许出现充电、放电或静置中的一种状态。为避免策略在收益目标下产生不可执行动作，项目还在目标函数中加入循环成本、短缺惩罚、弃光惩罚和终端 SOC 偏差惩罚。Stage12、Stage13 和 Stage15 的报告都记录了 SOC 越界、功率越限、同时充放电以及能量守恒误差等门禁结果，这些检查决定了策略能否进入后续比较。",
    "本文采用 24h rolling look-ahead 策略。输入为 Stage9 t+24h 预测结果、交付时刻对齐后的电价和负荷信号，目标是在满足 SOC 和功率约束的条件下进行价差利用，同时加入循环成本、短缺风险惩罚和终端 SOC 惩罚。相比离线阈值扫描，滚动优化策略收益较低，但更接近可执行调度逻辑。":
        "滚动策略的窗口长度设为 24h。每个调度时刻先读取 Stage9 的 t+24h 预测结果，再按交付时刻对齐电价和负荷信号，只执行当前一步充放电动作。目标函数同时考虑价差收益、循环成本、短缺风险和终端 SOC 偏差。它的收益低于一次性分位数扫描并不意外，因为后者更接近离线搜索；滚动策略的价值在于约束可审计，且更符合逐时更新的调度方式。",
    "本文预测模型主要使用 MAE、RMSE、nRMSE 和日间 nRMSE 评价。nRMSE 通过容量归一化，使不同模型的误差可比。储能调度主要评价总收益、相对无储能增量收益、充电量、放电量、等效循环次数、短缺量、弃光量、SOC 区间和约束门禁。系统实现部分则通过 API smoke、前端构建、E2E 测试和安全配置检查进行验证。":
        "评价指标按系统链路分成三组。预测模型侧记录 MAE、RMSE、nRMSE 和日间 nRMSE，其中 nRMSE 用装机容量归一化，便于比较不同模型。调度侧关注总收益、相对无储能增量收益、充放电量、等效循环次数、短缺量、弃光量、SOC 区间和约束门禁。系统侧不写成性能压测，而是使用 API smoke、前端构建、E2E 测试和安全配置检查来确认主要功能是否可演示。",
    "数据层负责接入 PVDAQ、NSRDB 和 OPSD 数据，完成清洗、标准化和特征工程。预测层负责训练 LightGBM、表格对比模型和深度学习模型，并固化主模型推理链路。调度层负责基于预测结果生成储能调度计划。治理层负责对策略收益、循环、短缺和 SOC 风险进行综合评分。服务层使用 FastAPI 提供统一接口。展示层使用 Vue 3 和 ECharts 展示系统结果。":
        "系统按照阶段产物组织，而不是把所有功能堆在一个脚本中。数据层输出可复用的清洗表和特征表；预测层在这些特征上训练 LightGBM、表格对比模型和深度学习模型，并把 Stage9 主模型推理结果固定下来；调度层读取预测、电价和负荷信号，生成储能动作及 SOC 序列；治理层再把收益、循环、短缺和 SOC 贴边风险转成可比较指标。FastAPI 提供统一读取入口，Vue 3 页面只消费接口返回的数据和报告摘要。",
    "数据模块负责读取配置文件、下载或加载原始数据、统一字段命名、处理缺失值并输出 parquet 和 CSV 文件。预测模块负责模型训练、模型对比、主模型推理和质量报告。调度模块负责固定阈值调度、分位数阈值扫描、滚动优化和配置敏感性分析。治理模块负责汇总不同策略指标，输出治理评分和风险标签。后端模块负责鉴权、数据读取和任务触发。前端模块负责页面展示、图表渲染和报告查看。":
        "从模块边界看，数据处理模块承担配置读取、原始数据接入、字段统一、缺失值处理以及 parquet/CSV 产物输出。预测模块围绕训练、对比、主模型推理和质量报告展开。调度模块接收预测产物后运行固定阈值、分位数扫描、滚动优化和配置敏感性分析。治理模块不重新计算调度动作，而是汇总策略指标并生成评分和风险标签。后端模块提供鉴权、数据读取和任务触发接口，前端模块则把指标、图表和 Markdown 报告组织成可演示页面。",
    "本文系统设计遵循四个原则。第一，数据可追溯。所有实验结果都应能定位到输入数据、配置文件和阶段报告。第二，输入边界清晰。严禁将目标时刻后验天气当作真实上线预测输入。第三，调度约束可审计。所有储能策略必须检查 SOC、功率、同时充放电和能量守恒。第四，结论不过度外推。OPSD 映射电价下的离线收益不能写成真实市场收益。":
        "系统设计主要受四个边界约束。其一，实验结果必须能追溯到输入数据、配置文件和阶段报告，否则不写入论文结论。其二，预测输入要区分 history_only 与目标时刻对齐字段，后者只能用于离线上限分析。其三，储能策略输出后必须检查 SOC、功率、同时充放电和能量守恒，不能只看收益。其四，OPSD 映射电价得到的是离线回放收益，不能被表述为 Colorado/PSCO 的真实市场结算结果。",
    "本文主线使用 PVDAQ system 10 光伏功率数据，地理位置为 39.7404、-105.1774。太阳资源和天气特征来自 NSRDB PSM 数据，负荷和电价画像来自 OPSD 数据。主实验时间范围为 2020 年至 2022 年。数据首先被统一为小时级时间序列，并转换为 UTC 时间。随后系统对重复时间戳、缺失值、异常值和字段单位进行检查。":
        "主实验围绕 PVDAQ system 10 展开，该站点位置为 39.7404、-105.1774。太阳资源字段来自 NSRDB PSM，负荷和电价画像来自 OPSD。由于这些数据源的时间口径和字段单位并不完全一致，项目先把 2020 年至 2022 年的数据统一到小时级 UTC 时间轴，再检查重复时间戳、缺失值、异常值和单位转换问题。只有通过这些检查后的表，才进入后续特征工程。",
    "三年数据链路的 Stage2 样本数为 25557，Stage3 特征工程后样本数为 25365，整体数据覆盖率达到建模要求。Stage3 删除的样本主要来自滞后窗口和未来标签构造。项目报告显示，三年数据相比早期较短时间窗口显著提升了 t+6h 和 t+24h 预测任务的稳定性。":
        "在三年数据链路中，Stage2 保留 25557 条样本，Stage3 完成滞后特征和未来标签构造后剩余 25365 条。减少的部分主要来自历史窗口不足或预测标签无法构造的边界时刻。与早期短时间窗口相比，2020-2022 数据覆盖了更多季节和天气变化，t+6h、t+24h 任务的训练和测试划分也更稳定。",
    "项目设定的替换规则要求候选模型在测试集 nRMSE 至少优于 LightGBM 0.0030，且日间 nRMSE 不变差。当前没有模型满足该规则，因此 LightGBM 被保留为主模型。":
        "主模型替换没有只看单个最优指标。项目门禁要求候选模型的测试集 nRMSE 至少比 LightGBM 低 0.0030，同时日间 nRMSE 不能变差。按这一规则检查后，当前候选模型都没有同时满足两项条件，因此 Stage9 继续采用 LightGBM 作为主模型。",
    "Stage9 将 LightGBM tuned history_only 模型固化为主推理链路。推理模块首先加载模型 bundle，然后校验特征列、容量配置、时间顺序、缺失值和无穷值，再执行批量推理，并将预测值裁剪到物理边界 [0, capacity_kw * 1.05]。Stage9 输出标准预测 CSV 和指标报告。":
        "Stage9 的作用是把训练阶段的 tuned LightGBM history_only 模型转成可复用推理产物。推理时，程序先读取模型 bundle，再逐项检查特征列、容量配置、时间顺序、缺失值和无穷值。批量预测完成后，结果被裁剪到 [0, capacity_kw * 1.05] 的物理范围内，并输出标准预测 CSV 与指标报告。后续调度模块读取的是这套标准产物，而不是临时实验结果。",
    "结果表明，Persistence 明显弱于机器学习和深度学习模型，说明 t+24h 光伏预测不是简单复制当前功率即可解决的问题。CNN-LSTM 在 target-aligned 离线上限条件下取得最优表现，但该输入不代表真实预测时刻可获得数据。在 history_only 条件下，CNN-LSTM 日间 nRMSE 略优于 LightGBM，但整体 nRMSE 未达到替代阈值。因此本文不将深度学习模型替换为主模型，而是将其作为预测算法验证和性能边界分析。":
        "从 Stage14B 结果看，Persistence 基线与机器学习、深度学习模型差距明显，说明 t+24h 光伏预测不能简单延用当前功率。CNN-LSTM 在 target-aligned 离线上限条件下表现最好，但这种输入包含目标时刻太阳资源信息，不能代表真实预测时刻的可获得数据。切换到 history_only 条件后，CNN-LSTM 的日间 nRMSE 略优于 LightGBM，整体 nRMSE 却没有达到替换门禁。因此，本文把深度学习实验用于说明模型边界，而不是替换主推理链路。",
    "本章完成了光伏功率预测模型设计与实现。通过多模型对比和质量门禁检查，本文选择 LightGBM tuned history_only 作为系统主模型，测试集 nRMSE 为 0.1225，日间 nRMSE 为 0.1689。深度学习模型完成了实现与对比，但在真实可获得输入条件下未达到替代主模型要求。该结果为后续储能调度提供了稳定预测输入。":
        "本章把预测模型从数据输入、训练对比到主模型固化进行了闭环处理。经过同口径实验和替换门禁检查，系统采用 LightGBM tuned history_only 作为主模型，其测试集 nRMSE 为 0.1225，日间 nRMSE 为 0.1689。深度学习模型已经完成实现和对照实验，但在真实可获得输入条件下没有触发替换规则。后续储能调度因此统一读取 Stage9 的预测产物，避免不同实验口径混用。",
    "Stage10 首先实现固定阈值调度策略。该策略在电价低于充电阈值时充电，在电价高于放电阈值时放电。固定阈值策略结构简单、可解释性强，适合作为基础基线。但实验发现，配置中的 discharge_price_threshold 为 45.0 EUR/MWh，而样本电价最大值约为 37.67 EUR/MWh，导致全周期几乎没有放电动作。该策略相对无储能的增量收益为 -0.0227 EUR。":
        "Stage10 用固定阈值策略作为最早的调度基线：电价低于充电阈值时充电，高于放电阈值时放电。这个规则便于检查，但也暴露出参数依赖问题。配置中的 discharge_price_threshold 为 45.0 EUR/MWh，而样本最大电价约为 37.67 EUR/MWh，放电条件在大部分时间无法触发。最终该策略相对无储能基线的增量收益为 -0.0227 EUR，说明阈值不能脱离样本分布设置。",
    "结果表明，适当增大容量并降低功率倍率有助于提高收益并控制循环强度。但推荐配置仍存在较高 SOC 贴边比例，说明当前调度目标函数仍需进一步优化。总体而言，储能配置和目标函数权衡对调度结果的影响大于继续无边界调参预测模型。":
        "Stage15 的敏感性结果显示，容量增大、功率倍率降低时，收益和循环强度之间的关系更容易控制。推荐配置仍然存在较高 SOC 贴边比例，这说明问题不只在储能硬件参数，也与目标函数中短缺、循环和终端 SOC 惩罚的权重有关。相比继续无边界地调参预测模型，明确储能配置与目标函数的取舍，对调度结果影响更直接。",
    "本文调度收益基于 OPSD 映射电价和离线 actual_kw 回放，不是真实 Colorado/PSCO 同区域市场收益。Stage15A 验证了 SPP WEIS RTBM LMP 作为真实市场扩展数据源的可行性，但 Xcel Energy-Colorado/PSCO 的 WEIS 财务绑定起点为 2023-04-01，晚于 2020-2022 主实验期。因此，WEIS 数据只能作为后续扩展验证，不能替换本文主实验收益口径。":
        "本文的收益口径需要单独限定：调度回放使用 OPSD 映射电价和离线 actual_kw，并不等同于 Colorado/PSCO 同区域真实市场结算。Stage15A 已检查 SPP WEIS RTBM LMP 的扩展可行性，但 Xcel Energy-Colorado/PSCO 与 WEIS 的财务绑定起点为 2023-04-01，晚于本文 2020-2022 主实验期。基于这一时间边界，WEIS 数据只能放在后续扩展验证中，不能替换当前主实验收益。",
    "本章完成储能调度与策略分析。固定阈值策略揭示了阈值与电价分布不匹配问题；分位数阈值策略给出离线上界；滚动优化策略形成可审计试点方案；策略治理和配置敏感性分析进一步说明了收益、循环、短缺和 SOC 风险之间的权衡。实验结果证明，储能调度质量不仅取决于预测精度，更取决于策略参数、目标函数和储能配置。":
        "本章从基线、上界和可执行策略三个层次分析储能调度。固定阈值策略暴露了阈值与电价分布不匹配的问题；分位数阈值扫描提供离线对照；滚动优化策略虽然收益较低，但具备 SOC 和功率约束审计条件。Stage13 的治理评分和 Stage15 的配置敏感性进一步说明，收益、循环、短缺和 SOC 风险会相互牵制。由此可见，调度质量不能只归因于预测误差，策略参数、目标函数和储能配置同样决定最终结果。",
    "系统后端采用 Python 和 FastAPI 实现。核心功能包括用户鉴权、配置读取、模型指标查询、预测结果查询、调度指标查询、治理评分查询、敏感性结果查询、报告列表读取和后端任务提交。后端将阶段产物统一封装为 API，前端不直接读取底层 CSV 或 Markdown 文件，而是通过接口获取数据。":
        "后端使用 Python 与 FastAPI 组织接口，主要面向两类数据：一类是预测、调度、治理和敏感性分析形成的阶段产物，另一类是报告归档和任务触发所需的运行信息。用户登录后，前端通过接口读取模型指标、预测结果、调度指标、治理评分和报告列表。这样处理后，CSV、Markdown 等底层文件不会直接暴露给页面，接口层也能集中处理鉴权和数据格式转换。",
    "后端安全方面，系统在生产环境中要求显式配置 NES_APP_ENV、NES_JWT_SECRET、NES_USERS_JSON 和 NES_CORS_ORIGINS。若生产环境使用默认 secret、默认用户或宽松 CORS，后端会拒绝启动。这一设计避免演示账号和开发配置被误用于生产环境。":
        "安全配置主要针对演示环境和生产环境混用的风险。后端在生产模式下检查 NES_APP_ENV、NES_JWT_SECRET、NES_USERS_JSON 和 NES_CORS_ORIGINS；如果仍使用默认 secret、默认用户或过宽的 CORS 配置，应用会在启动阶段中止。本文没有把该机制写成完整安全体系，而是将其作为防止开发配置误带入生产环境的基本保护。",
    "前端采用 Vue 3、Vite、Element Plus 和 ECharts 实现。主要页面包括系统总览、模型评估、策略收益、配置治理、任务运维和报告归档。系统总览展示关键指标和运行状态；模型评估展示 LightGBM、表格模型和深度学习模型指标；策略收益展示储能充放电、SOC 和收益曲线；配置治理展示策略评分和敏感性分析；报告归档用于查看各阶段 Markdown 报告。":
        "前端由 Vue 3 和 Vite 构建，组件层使用 Element Plus，图表展示使用 ECharts。页面按照答辩演示路径组织：系统总览先给出关键指标和运行状态；模型评估页对比 LightGBM、表格模型和深度学习模型；策略收益页展示充放电、SOC 和收益曲线；配置治理页查看策略评分与敏感性结果；报告归档页读取各阶段 Markdown 报告。这样的页面划分对应项目阶段产物，便于从实验结果追溯到系统功能。",
    "系统测试包括后端静态检查、API smoke 测试、前端构建和 Playwright E2E 测试。后端通过 py_compile 检查关键模块语法，通过 API smoke 验证登录和核心接口返回。前端通过 npm run lint、npm run build 和 npm run test:e2e 验证构建和交互流程。":
        "测试工作按后端、接口和前端三条线执行。后端先用 py_compile 检查关键模块是否存在语法错误，再通过 API smoke 覆盖登录和核心接口返回。前端侧执行 npm run lint 和 npm run build，确认代码规范与生产构建可通过；交互流程则交给 Playwright E2E 测试。本文没有声称完成压力测试，只把这些结果作为原型系统可演示性的依据。",
    "E2E 测试覆盖登录页不展示明文 demo 凭据、admin 登录、核心路由导航、移动端横向溢出、Markdown 报告 XSS 净化和 guest 提交任务权限失败可见错误。当前 E2E 结果为 4 个测试通过。":
        "Playwright E2E 用例覆盖了几类容易在演示中暴露的问题：登录页不能直接展示明文 demo 凭据，admin 用户需要能够进入核心页面，主要路由应能正常导航，移动端不能出现明显横向溢出。报告页还检查 Markdown 内容的 XSS 净化，guest 用户提交任务失败时需要看到明确错误提示。当前记录中 4 个 E2E 用例通过。",
    "当前系统已经达到毕业设计演示和继续部署加固的原型标准，但仍不能直接宣称完成生产部署。生产部署还需要在目标机器上配置环境变量、反向代理、前端静态资源服务、后端进程管理和 HTTPS，并重新执行 API smoke 与 E2E 验收。":
        "从现有测试结果看，系统可以支撑毕业设计演示，也具备继续部署加固的基础。但本文不把它表述为已经生产上线。若要进入真实部署，还需要在目标机器上重新配置环境变量、反向代理、前端静态资源服务、后端进程管理和 HTTPS，并在部署环境中再次执行 API smoke 与 E2E 验收。",
    "第二，完成了多类光伏功率预测模型对比。实验表明，在严格 history_only 输入条件下，LightGBM tuned 模型在 t+24h 预测任务上取得测试集 nRMSE 0.1225 和日间 nRMSE 0.1689，并被固化为系统主模型。深度学习模型完成了实现和对比，但在真实可获得输入条件下尚未达到替代主模型的工程阈值。":
        "第二，本文完成了多类光伏预测模型的同口径比较。严格采用 history_only 输入后，LightGBM tuned 在 t+24h 任务上的测试集 nRMSE 为 0.1225，日间 nRMSE 为 0.1689，因此被写入 Stage9 主推理链路。深度学习模型已经实现并参与对比，但在真实可获得输入条件下没有通过主模型替换门禁。",
    "第四，完成了后端 API 和前端展示系统。系统能够展示预测指标、调度结果、治理评分、敏感性分析和阶段报告，并通过构建、接口和 E2E 测试验证主要功能。":
        "第四，本文把阶段产物接入后端 API 和前端页面。用户可以在页面上查看预测指标、调度曲线、治理评分、敏感性分析和阶段报告；这些页面的基础功能已经通过前端构建、接口 smoke 和 E2E 用例检查。",
    "本文仍存在以下不足。第一，主实验收益使用 OPSD 映射电价，不是真实 Colorado/PSCO 同区域结算价格。后续可基于 2023-04-01 后 SPP WEIS RTBM LMP 数据开展扩展验证。第二，Stage7 真实 forecast-cycle 天气替代实验尚未达到上线门槛，未来可继续接入 HRRR 原生 forecast issue/lead-time 数据。第三，当前滚动优化策略仍较简化，可进一步引入更严格的寿命模型、需求响应约束和多目标优化。第四，前端系统已满足演示要求，但在生产部署、包体优化和图表精修方面仍有改进空间。":
        "本文的不足主要来自数据边界和工程深度两方面。收益评估仍使用 OPSD 映射电价，不能替代 Colorado/PSCO 同区域结算价格；Stage15A 虽然确认 2023-04-01 之后的 SPP WEIS RTBM LMP 可用于扩展，但它无法覆盖本文主实验期。天气输入方面，Stage7 的真实 forecast-cycle 替代实验还没有达到上线门槛，后续需要继续接入 HRRR 原生 forecast issue 和 lead-time 数据。调度算法方面，当前滚动优化仍偏简化，寿命模型、需求响应约束和多目标优化可以继续补充。前端已满足演示，但生产部署、包体优化和图表细节仍需单独加固。",
    "总体来看，本文完成了新能源储能侧从光伏预测到调度优化再到系统展示的完整原型，为后续接入真实市场数据、增强 forecast-cycle 天气输入和改进储能优化算法提供了基础。":
        "综合来看，本文已经把光伏预测、储能调度和系统展示串成可运行原型。后续工作可以沿三条线继续推进：用真实市场数据替换当前 OPSD 映射电价，增强 forecast-cycle 天气输入的可用性，并在滚动优化中加入更细的寿命和多目标约束。",
    "本课题从数据处理、模型训练、储能调度到系统展示，涉及电力系统、新能源预测、机器学习、后端服务和前端可视化等多个方向。在课题推进过程中，指导教师在研究方向确定、技术路线选择和论文结构梳理方面给予了重要帮助，使本人能够在复杂任务中逐步明确系统边界和实现重点。":
        "本课题从最初的数据整理推进到模型训练、储能调度和系统展示，过程中多次需要在算法效果、数据可获得性和工程实现之间取舍。指导教师在研究方向确定、技术路线筛选和论文结构调整方面给了我很多具体建议，尤其是在区分离线上限、真实可获得输入和系统演示边界时，帮助我把课题范围逐步收敛下来。",
    "同时，感谢学院在毕业设计过程中的组织与支持，使本人能够围绕实际工程问题完成较完整的系统实现。通过本课题的研究与开发，本人加深了对新能源功率预测、储能优化调度和软件工程闭环的理解，也提升了独立分析问题、定位风险和整理实验结论的能力。":
        "同时感谢学院在毕业设计各阶段提供的组织和支持。完成这个课题后，我对新能源功率预测、储能调度约束和前后端系统衔接有了更直接的认识，也在反复检查数据口径、模型边界和实验结论的过程中，提高了独立分析问题和整理工程材料的能力。",
}


PROCESS_WORDS = ["风险点", "改写段落", "检测报告", "提示词"]


def set_paragraph_text(paragraph, text: str) -> None:
    """Replace paragraph text while preserving the original paragraph style and basic run font."""
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def rewrite_docx(source: Path, output: Path) -> tuple[int, list[str]]:
    """Copy the thesis and rewrite all configured high/medium-risk paragraphs."""
    shutil.copy2(source, output)
    document = Document(str(output))
    replaced = 0
    missed: list[str] = []
    remaining = {compact(key): (key, value) for key, value in REPLACEMENTS.items()}

    for paragraph in document.paragraphs:
        normalized = compact(paragraph.text)
        if normalized in remaining:
            _, replacement = remaining.pop(normalized)
            set_paragraph_text(paragraph, replacement)
            replaced += 1

    missed.extend(original[:80] for original, _ in remaining.values())
    document.save(str(output))
    return replaced, missed


def validate_output(output: Path) -> dict[str, bool]:
    """Run lightweight checks that catch leakage of process text and broken DOCX output."""
    document = Document(str(output))
    full_text = "\n".join(p.text for p in document.paragraphs)
    return {
        "DOCX 可读取": len(document.paragraphs) > 50,
        "无过程提示词泄露": not any(word in full_text for word in PROCESS_WORDS),
        "保留摘要": "摘 要" in full_text,
        "保留参考文献": "参考文献" in full_text,
        "保留核心指标": "nRMSE" in full_text and "0.1225" in full_text and "0.1689" in full_text,
        "保留关键数据源": all(token in full_text for token in ["PVDAQ", "NSRDB", "OPSD"]),
    }


def write_report(
    report_path: Path,
    source: Path,
    output: Path,
    html_report: Path,
    hotspots: list[tuple[str, str]],
    replaced: int,
    missed: list[str],
    checks: dict[str, bool],
) -> None:
    """Write a separate rework report so no process notes enter the thesis itself."""
    hotspot_counter = Counter(cls for cls, _ in hotspots)
    lines = [
        "# 降 AIGC 修改报告",
        "",
        f"- 源文件：`{source}`",
        f"- 输出文件：`{output}`",
        f"- 检测报告：`{html_report}`",
        f"- 修改时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        "- PaperYY 原报告：疑似率 41%，高度疑似 40.1%，中度疑似 3.5%，低度疑似 27.9%。",
        f"- 报告解析片段：high={hotspot_counter.get('high', 0)}，medium={hotspot_counter.get('medium', 0)}。",
        f"- 本轮替换段落数：{replaced}",
        "",
        "## 改写策略",
        "- 保留项目事实、模型名、阶段编号、指标、引用编号和实验边界。",
        "- 将模板化的“负责、完成、表明、验证了、主要包括”改为数据链路、阶段产物和约束检查描述。",
        "- 优先改写摘要、绪论、理论基础、总体设计、实验分析、系统测试、总结与致谢中的 high/medium 段落。",
        "- 未改低度疑似且技术事实密集的段落，避免引入事实错误。",
        "",
        "## 自动检查",
    ]
    for name, passed in checks.items():
        lines.append(f"- {'通过' if passed else '未通过'}：{name}")
    lines.extend(["", "## 未匹配段落"])
    if missed:
        lines.extend(f"- {item}" for item in missed)
    else:
        lines.append("- 无。")
    lines.extend(
        [
            "",
            "## 复检建议",
            "- 使用输出文件重新提交 PaperYY 检测。",
            "- 第二轮只处理新报告中仍为 high/medium 的段落，不建议全文盲改。",
            "- 如果目录页码变化，打开 Word 后执行“更新整个目录”。",
            "",
            "## Pitfall",
            "- AIGC 检测结果具有模型和版本差异，本轮只能降低模板化表达风险，不能承诺固定比例。",
            "",
            "## 阶段进度",
            "- 本阶段目标：完成报告驱动的第一轮 high/medium 段落改写。",
            "- 完成情况：已生成降 AIGC 版 DOCX 与独立修改报告。",
            "- 下一阶段可行性：高。复检后可按剩余高风险段落做第二轮定点改写。",
        ]
    )
    report_path.write_text("\n".join(lines), encoding="utf-8")


def main() -> None:
    source = find_source_docx()
    html_report = find_report_html()
    output = source.with_name(f"{source.stem}_降AIGC版{source.suffix}")
    report_path = source.with_name(f"{source.stem}_降AIGC修改报告.md")

    hotspots = parse_hotspots(html_report)
    replaced, missed = rewrite_docx(source, output)
    checks = validate_output(output)
    write_report(report_path, source, output, html_report, hotspots, replaced, missed, checks)

    print(f"SOURCE={source}")
    print(f"OUTPUT={output}")
    print(f"REPORT={report_path}")
    print(f"HOTSPOTS={Counter(cls for cls, _ in hotspots)}")
    print(f"REPLACED={replaced}")
    print(f"MISSED={len(missed)}")
    for name, passed in checks.items():
        print(f"CHECK {name}: {'PASS' if passed else 'FAIL'}")
    if missed or not all(checks.values()):
        raise SystemExit(2)


if __name__ == "__main__":
    main()
